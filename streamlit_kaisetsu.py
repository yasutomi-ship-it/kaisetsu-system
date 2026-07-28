"""
薬剤師国家試験 解説生成システム（Streamlit版）
ブラウザで完結。ドラッグ＆ドロップ→生成→ダウンロード。
"""

import streamlit as st
import json
import re
import copy
import io
import urllib.request
import urllib.error
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt

# ──────────────────────────────────────────────
# フォーマットエンジン
# ──────────────────────────────────────────────
TNR      = 'Times New Roman'
CENTURY  = 'Century'
GOTHIC   = 'ＭＳ ゴシック'
MINCHO   = 'ＭＳ 明朝'
HIRAGINO = 'ヒラギノ角ゴ ProN W3'
MARKUP   = re.compile(r'\{\{([^}]+)\}\}')
NUMBERS  = ['１', '２', '３', '４', '５']


_DASH_CHARS = ('-', '−', '‐', '‑', '–', '—',
               '－')  # 半角ハイフン/数学マイナス/各種ハイフン・ダッシュ/全角


def normalize_dashes(text):
    """連結記号・マイナスを全角ハイフンマイナス「－」（U+FF0D）に統一する。
    半角ハイフン・数学マイナス(U+2212)・各種ダッシュを吸収し、
    フォントが周囲と揃うようにする。"""
    for ch in _DASH_CHARS:
        text = text.replace(ch, '－')
    return text


# {{...}}（特殊フォント用）、_{...}（下付き）、^{...}（上付き）を1つのトークナイザで処理
TOKEN_RE = re.compile(r'(\{\{[^}]+\}\})|([_^])\{([^}]*)\}')


def _fw_signs(text):
    """上付き・下付き内のプラス・マイナスを全角（＋＝U+FF0B／－＝U+FF0D）にする。
    全角記号はMS明朝（eastAsiaフォント）で描画され、社内フォーマットと字体が揃う。"""
    return text.replace('-', '－').replace('−', '－').replace('+', '＋')


def auto_scientific(text):
    """上付き・下付きが markup 化されていない一般的な理系表記を _{ }/^{ } に自動変換する
    （AIが markup を付け忘れた場合の保険）。"""
    # 括弧付き指数 e^(-kt) → e^{-kt}
    text = re.sub(r'\^\(([^)]*)\)', r'^{\1}', text)
    # 10のべき乗：10－5 / 10-5 → 10^{-5}
    text = re.sub(r'10[－\-−]([0-9]+)', r'10^{-\1}', text)
    # 半減期・初期濃度など定番の下付き（薬学分野で一意）
    text = text.replace('t1/2', 't_{1/2}')
    text = re.sub(r'([AC])0(?![0-9])', r'\1_{0}', text)  # A0, C0
    return text


def _expand_curly(c):
    """{{...}} トークン1つ分を runs に展開する。"""
    if   c == '-':            return [{'text': '－'}]
    elif c == 'mu':           return [{'text': 'µ', 'font': CENTURY, 'italic': True}]
    elif c.startswith('sup:'): return [{'text': _fw_signs(c[4:]), 'sup': True, 'font': MINCHO}]
    elif c.startswith('sub:'): return [{'text': _fw_signs(c[4:]), 'sub': True, 'font': MINCHO}]
    elif c.startswith('CL:'):  return [{'text':'CL','font':TNR,'italic':True},{'text':_fw_signs(c[3:]),'sub':True}]
    elif c == 'CL':           return [{'text':'CL','font':TNR,'italic':True}]
    elif c.startswith('f:'):   return [{'text':'f','font':CENTURY,'italic':True},{'text':_fw_signs(c[2:]),'sub':True}]
    elif c == 'f':            return [{'text':'f','font':CENTURY,'italic':True}]
    elif c.startswith('K:'):   return [{'text':'K','font':CENTURY,'italic':True},{'text':_fw_signs(c[2:]),'sub':True}]
    elif c.startswith('t:'):   return [{'text':'t','font':CENTURY,'italic':True},{'text':_fw_signs(c[2:]),'sub':True}]
    elif 'Vd' in c:
        r = [{'text':'Vd','font':CENTURY,'italic':True}]
        if ':' in c: r.append({'text':_fw_signs(c.split(':',1)[1]),'sub':True})
        return r
    return [{'text': c}]


def parse_markup(text):
    text = auto_scientific(text)
    runs, pos = [], 0
    for m in TOKEN_RE.finditer(text):
        if m.start() > pos:
            runs.append({'text': normalize_dashes(text[pos:m.start()])})
        if m.group(1):                       # {{...}}
            runs += _expand_curly(m.group(1)[2:-2])
        else:                                # _{...} / ^{...}
            kind, inner = m.group(2), m.group(3)
            runs.append({'text': _fw_signs(inner), 'sub': kind == '_',
                         'sup': kind == '^', 'font': MINCHO})
        pos = m.end()
    if pos < len(text):
        runs.append({'text': normalize_dashes(text[pos:])})
    return runs


def split_bullets(text):
    """「●」区切りの列挙を各行に分割する（前文で複数種を提示する場合に見やすくする）。"""
    if '●' not in text:
        return [text]
    parts = re.split(r'(?=●)', text)
    return [p.strip() for p in parts if p.strip()]


def make_run(text, font=None, italic=False, bold=None, sub=False, sup=False):
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if font:
        rf = OxmlElement('w:rFonts')
        for a in ('w:ascii','w:hAnsi','w:eastAsia','w:cs'): rf.set(qn(a), font)
        rPr.append(rf)
    if bold is True:  rPr.append(OxmlElement('w:b'))
    elif bold is False:
        b = OxmlElement('w:b'); b.set(qn('w:val'),'0'); rPr.append(b)
    if italic: rPr.append(OxmlElement('w:i')); rPr.append(OxmlElement('w:iCs'))
    if sub or sup:
        va = OxmlElement('w:vertAlign')
        va.set(qn('w:val'), 'subscript' if sub else 'superscript')
        rPr.append(va)
    if len(rPr): r.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    if text and (text[0]==' ' or text[-1]==' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space','preserve')
    r.append(t)
    return r


def add_runs(p_elem, specs):
    for s in specs:
        p_elem.append(make_run(s.get('text',''), font=s.get('font'),
            italic=s.get('italic',False), bold=s.get('bold'),
            sub=s.get('sub',False), sup=s.get('sup',False)))


def clear_runs(para):
    for r in para._p.findall(qn('w:r')): para._p.remove(r)


EXP_PAT = re.compile(r'^[１２３４５]　(誤|正)：')
HEAD_PAT = re.compile(r'^問\s*(\d+)\s*(?:（([^）]*)）)?')


def _is_answer_line(text):
    t = text.strip()
    return t.startswith('問') and '解答' in t


def segment_questions(doc):
    """テンプレートを小問ブロックに分割する。
    連問（問196以降）は複数の解答行を持つため、解答行を区切りとしてブロック化する。
    各ブロック: 問番号・科目・前文プレースホルダ・選択肢解説プレースホルダ・解答行。
    単問の場合は要素1つのリストを返す。"""
    paras = doc.paragraphs
    ans_idxs = [i for i, p in enumerate(paras) if _is_answer_line(p.text)]
    blocks = []
    prev = -1
    for ai in ans_idxs:
        exps = [i for i in range(prev + 1, ai) if EXP_PAT.match(paras[i].text)]
        if not exps:
            prev = ai
            continue
        first = exps[0]
        # ブロック内の問番号ヘッダー（問NNN（科目））
        header_idx, qnum, subject = None, None, None
        for i in range(first - 1, prev, -1):
            t = paras[i].text.strip()
            if t.startswith('問') and '解答' not in t:
                m = HEAD_PAT.match(t)
                if m:
                    header_idx, qnum, subject = i, m.group(1), m.group(2)
                    break
        # 解答行からも問番号を補完
        if qnum is None:
            m = HEAD_PAT.match(paras[ai].text.strip())
            if m:
                qnum = m.group(1)
        # 前文プレースホルダ（選択肢解説の直前・非空・番号始まりでない段落）
        pre_idx = None
        for i in range(first - 1, (header_idx if header_idx is not None else prev), -1):
            t = paras[i].text.strip()
            if not t:
                continue
            if EXP_PAT.match(paras[i].text):
                continue
            if t[0] in '１２３４５' or t.startswith('問') or '解答' in t:
                break  # 実際の選択肢・見出しに到達 → 前文なし
            pre_idx = i
            break
        blocks.append({
            'question_num': qnum, 'subject': subject,
            'header_para_idx': header_idx, 'preamble_para_idx': pre_idx,
            'explanation_para_idxs': exps, 'answer_para_idx': ai,
        })
        prev = ai
    return blocks


def extract_info(doc):
    """後方互換: 最初の小問ブロックを単問形式で返す。"""
    blocks = segment_questions(doc)
    if not blocks:
        return {'question_num': None, 'preamble_para_idx': None,
                'explanation_para_idxs': [], 'answer_para_idx': None}
    b = blocks[0]
    return {'question_num': b['question_num'],
            'preamble_para_idx': b['preamble_para_idx'],
            'explanation_para_idxs': b['explanation_para_idxs'],
            'answer_para_idx': b['answer_para_idx']}


def _write_header(p, qnum, subject):
    """問番号ヘッダーを問題側と統一（問＋番号をボールド、ＭＳゴシック）。"""
    clear_runs(p)
    p._p.append(make_run('問', font=GOTHIC, bold=True))
    p._p.append(make_run(str(qnum or '000'), font=GOTHIC, bold=True))
    if subject:
        p._p.append(make_run('（' + subject + '）', font=GOTHIC))


def write_block(doc, block, expl):
    """1つの小問ブロックに前文・選択肢解説・解答を書き込む（削除はしない）。"""
    paras = doc.paragraphs
    # ヘッダー（「問N（科目）」形式の短い見出しのみ、問番号をボールド化して統一）
    # 問題文の設問行（長文）を誤って書き換えないよう文字数でガードする
    if block.get('header_para_idx') is not None:
        hp = paras[block['header_para_idx']]
        if len(hp.text.strip()) <= 12 and re.match(r'^問\d+', hp.text.strip()):
            _write_header(hp, block.get('question_num'), block.get('subject'))

    # 選択肢解説
    for idx, item in zip(block['explanation_para_idxs'], expl.get('選択肢解説', [])):
        p = paras[idx]
        clear_runs(p)
        p._p.append(make_run('', font=GOTHIC, bold=True))
        p._p.append(make_run(item.get('番号', ''), font=GOTHIC))
        p._p.append(make_run('　', font=HIRAGINO))
        p._p.append(make_run(item.get('正誤', '誤') + '：'))
        add_runs(p._p, parse_markup(item.get('内容', '')))

    # 解答行（右揃え・全体ボールド）
    p_ans = paras[block['answer_para_idx']]
    clear_runs(p_ans)
    p_ans.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    q = block.get('question_num', '000')
    for seg in ['問', str(q), '　', '解答　']:
        p_ans._p.append(make_run(seg, font=GOTHIC, bold=True))
    for i, ans in enumerate(expl.get('解答', [])):
        p_ans._p.append(make_run(ans, font=GOTHIC, bold=True))
        if i < len(expl['解答']) - 1:
            p_ans._p.append(make_run('、', font=GOTHIC, bold=True))

    # 前文
    if block.get('preamble_para_idx') is not None:
        p_pre = paras[block['preamble_para_idx']]
        clear_runs(p_pre)
        前文 = expl.get('前文', '')
        add_runs(p_pre._p, parse_markup(前文) if 前文 else [{'text': ''}])
        pPr = p_pre._p.find(qn('w:pPr'))
        last = p_pre._p
        for line in expl.get('前文_追加行', []):
            np = OxmlElement('w:p')
            if pPr is not None:
                np.append(copy.deepcopy(pPr))
            last.addnext(np)
            add_runs(np, parse_markup(line))
            last = np


def write_all(doc, blocks, expls):
    """全ブロックを書き込み、最後の解答行より後ろのテンプレート雛形を削除する。"""
    # 末尾雛形の削除（最後の解答行以降）
    last_ans = doc.paragraphs[blocks[-1]['answer_para_idx']]._p
    del_elems, found = [], False
    for p in doc.paragraphs:
        if found:
            del_elems.append(p._p)
        if p._p is last_ans:
            found = True
    # 書き込みは段落インデックスに依存するため、削除より前に実施
    for block, expl in zip(blocks, expls):
        write_block(doc, block, expl)
    for e in del_elems:
        if e.getparent() is not None:
            e.getparent().remove(e)


def write_to_doc(doc, info, expl):
    """後方互換: 単問1つ分を書き込む。"""
    block = {'question_num': info.get('question_num'), 'subject': None,
             'header_para_idx': None,
             'preamble_para_idx': info.get('preamble_para_idx'),
             'explanation_para_idxs': info['explanation_para_idxs'],
             'answer_para_idx': info['answer_para_idx']}
    write_all(doc, [block], [expl])


# ──────────────────────────────────────────────
# 社内フォーマット対応：問題本文から小問を自動検出し、
# 解説を実番号で再生成する（雛形ライブラリはそのまま入っていてよい）
# ──────────────────────────────────────────────
NEG_MARKERS = ('誤っているのはどれか', '正しくないのはどれか', '適切でないのはどれか',
               '適切でないもの', '含まれないのはどれか', '該当しないのはどれか',
               'ないのはどれか')
# 雛形（プレースホルダ）セクションの開始を示す目印
TEMPLATE_MARKERS = ('解説（一括', 'ひな型', '雛形', '否定文用', '否定文ひな型',
                    '連問の場合')


def _set_ind(p, **kw):
    pPr = p._element.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind'); pPr.append(ind)
    for k, v in kw.items():
        ind.set(qn('w:' + k), str(v))


def parse_question_blocks(doc):
    """問題本文から実際の小問を抽出する。
    返り値: (subs, boundary)
      subs = [{'num','subject','n_choices','neg'}...]（問題文に現れる順）
      boundary = 雛形（プレースホルダ）が始まる段落index。これ以降は解説再生成時に削除する。
    """
    paras = doc.paragraphs
    boundary = len(paras)
    for i, p in enumerate(paras):
        t = p.text.strip()
        if not t:
            continue
        if any(mk in t for mk in TEMPLATE_MARKERS) or EXP_PAT.match(t) or _is_answer_line(t):
            boundary = i
            break

    # 小問ヘッダー（boundaryより前の「問\d+…」。範囲タイトル「問A〜B」は除外）
    header_idxs = []
    for i in range(boundary):
        t = paras[i].text.strip()
        if re.match(r'^問\d+', t) and '〜' not in t and '～' not in t and '解答' not in t:
            header_idxs.append(i)

    subs = []
    for k, hi in enumerate(header_idxs):
        t = paras[hi].text.strip()
        num = re.match(r'^問(\d+)', t).group(1)
        sm = re.match(r'^問\d+（([^）]*)）', t)
        subject = sm.group(1) if sm else None
        end = header_idxs[k + 1] if k + 1 < len(header_idxs) else boundary
        seg_text = '\n'.join(paras[j].text for j in range(hi, end))
        n_choices = len(re.findall(r'[１２３４５６]　', seg_text))
        if n_choices == 0:
            n_choices = 5
        neg = any(mk in seg_text for mk in NEG_MARKERS)
        subs.append({'num': num, 'subject': subject,
                     'n_choices': n_choices, 'neg': neg})
    return subs, boundary


def delete_from(doc, boundary):
    """boundary以降の段落をすべて削除（雛形ライブラリの除去）。"""
    for p in list(doc.paragraphs)[boundary:]:
        if p._p.getparent() is not None:
            p._p.getparent().remove(p._p)


def smiles_to_png(smiles, size=(260, 200)):
    """SMILES から2D構造式のPNGバイト列を生成する。
    RDKitが無い／SMILESが不正な場合は None を返す（機能はグレースフルにスキップ）。"""
    try:
        from rdkit import Chem
        from rdkit.Chem import Draw
        from rdkit.Chem.Draw import rdMolDraw2D
        mol = Chem.MolFromSmiles(smiles)
        if mol is None:
            return None
        drawer = rdMolDraw2D.MolDraw2DCairo(size[0], size[1])
        opt = drawer.drawOptions()
        opt.bondLineWidth = 2
        drawer.DrawMolecule(mol)
        drawer.FinishDrawing()
        return io.BytesIO(drawer.GetDrawingText())
    except Exception:
        # Cairo が無い環境向けフォールバック
        try:
            from rdkit import Chem
            from rdkit.Chem import Draw
            mol = Chem.MolFromSmiles(smiles)
            if mol is None:
                return None
            img = Draw.MolToImage(mol, size=size)
            bio = io.BytesIO()
            img.save(bio, format='PNG')
            bio.seek(0)
            return bio
        except Exception:
            return None


def _clear_cell_borders(table):
    """表の罫線を消して図版らしく見せる。"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'none')
        borders.append(e)
    tblPr.append(borders)


def insert_structure_figure(doc, structures, per_row=3):
    """構造式リスト（[{'ラベル','smiles'}...]）を、ラベル付き画像の図版として挿入する。
    RDKit未導入・不正SMILESの項目は自動的にスキップする。"""
    rendered = []
    for s in structures or []:
        if not isinstance(s, dict):
            continue
        smi = s.get('smiles') or s.get('SMILES')
        if not smi:
            continue
        png = smiles_to_png(smi)
        if png is not None:
            rendered.append((s.get('ラベル', ''), png))
    if not rendered:
        return
    n = len(rendered)
    cols = min(per_row, n)
    rows = (n + cols - 1) // cols
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _clear_cell_borders(table)
    for idx, (label, png) in enumerate(rendered):
        r, c = divmod(idx, cols)
        cell = table.cell(r, c)
        pimg = cell.paragraphs[0]
        pimg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pimg.add_run().add_picture(png, width=Inches(1.7))
        pcap = cell.add_paragraph()
        pcap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = pcap.add_run(label)
        rr.font.size = Pt(9)


def append_explanation(doc, sub, expl, add_header=True):
    """問題本文の後ろに、実番号で書式付きの解説ブロックを追加する。
    書式は社内フォーマット（前文1字下げ／選択肢ぶら下げ／解答は右揃え太字）に準拠。"""
    doc.add_paragraph()  # 区切りの空行
    if add_header:
        p = doc.add_paragraph()
        p._p.append(make_run('問', font=GOTHIC, bold=True))
        p._p.append(make_run(str(sub.get('num', '000')), font=GOTHIC, bold=True))
        if sub.get('subject'):
            p._p.append(make_run('（' + sub['subject'] + '）', font=GOTHIC))

    # 前文（あれば）。「●」区切りの列挙は各行に分割して見やすくする。
    pre_lines = []
    if expl.get('前文'):
        pre_lines += split_bullets(expl['前文'])
    for line in expl.get('前文_追加行', []):
        if line:
            pre_lines += split_bullets(line)
    for line in pre_lines:
        p = doc.add_paragraph()
        add_runs(p._p, parse_markup(line))
        _set_ind(p, firstLineChars=100, firstLine=180)
    if pre_lines:
        doc.add_paragraph()  # 前文と選択肢の間の空行

    # 構造式（あれば）：ラベル付き図版として前文の後に挿入
    if expl.get('構造式'):
        insert_structure_figure(doc, expl['構造式'])
        doc.add_paragraph()

    # 選択肢解説（ぶら下げインデント）
    for item in expl.get('選択肢解説', []):
        p = doc.add_paragraph()
        p._p.append(make_run(item.get('番号', ''), font=GOTHIC))
        p._p.append(make_run('　', font=HIRAGINO))
        p._p.append(make_run(item.get('正誤', '誤') + '：'))
        add_runs(p._p, parse_markup(item.get('内容', '')))
        _set_ind(p, leftChars=100, left=720, hangingChars=300, hanging=540)

    doc.add_paragraph()  # 選択肢と解答の間の空行

    # 解答（右揃え・太字）
    p = doc.add_paragraph()
    for seg in ['問', str(sub.get('num', '000')), '　解答　']:
        p._p.append(make_run(seg, font=GOTHIC, bold=True))
    ans = expl.get('解答', [])
    for i, a in enumerate(ans):
        p._p.append(make_run(a, font=GOTHIC, bold=True))
        if i < len(ans) - 1:
            p._p.append(make_run('、', font=GOTHIC, bold=True))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def regenerate_explanations(doc, subs, boundary, expls):
    """雛形を削除し、問題本文の後ろに実番号で解説を再生成する。
    連問（複数小問）のときだけ各解説ブロックに問番号ヘッダーを付ける
    （単問は従来フォーマットどおりヘッダーなし）。"""
    delete_from(doc, boundary)
    add_header = len(subs) > 1
    for sub, expl in zip(subs, expls):
        append_explanation(doc, sub, expl, add_header=add_header)


# ──────────────────────────────────────────────
# PDF（画像認識）モード：Word発PDFを読み、house形式のdocxを新規生成する
# ──────────────────────────────────────────────
def new_house_doc():
    """社内フォーマットの既定スタイル（本文=Century/ＭＳ明朝 9pt）でDocumentを作る。"""
    doc = Document()
    stl = doc.styles['Normal']
    stl.font.name = 'Century'
    stl.font.size = Pt(9)
    rpr = stl.element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    rf.set(qn('w:ascii'), 'Century')
    rf.set(qn('w:hAnsi'), 'Century')
    rf.set(qn('w:eastAsia'), MINCHO)
    rf.set(qn('w:cs'), 'Times New Roman')
    return doc


def append_pdf_item(doc, item):
    """PDF小問1つ分を、設問＋選択肢＋解説（構造式込み）としてdocxに追加する。"""
    sub = {'num': str(item.get('問番号', '')), 'subject': item.get('科目')}
    # 見出し（問番号＋科目、問番号をボールド）
    p = doc.add_paragraph()
    p._p.append(make_run('問', font=GOTHIC, bold=True))
    p._p.append(make_run(sub['num'], font=GOTHIC, bold=True))
    if sub['subject']:
        p._p.append(make_run('（' + sub['subject'] + '）', font=GOTHIC))
    # 設問文
    stem = item.get('設問', '')
    if stem:
        ps = doc.add_paragraph()
        add_runs(ps._p, parse_markup(stem))
        _set_ind(ps, firstLineChars=100, firstLine=180)
    # 選択肢
    for i, ch in enumerate(item.get('選択肢', []) or []):
        pc = doc.add_paragraph()
        pc._p.append(make_run(NUMBERS[i] if i < len(NUMBERS) else str(i + 1), font=GOTHIC))
        pc._p.append(make_run('　', font=HIRAGINO))
        add_runs(pc._p, parse_markup(ch))
        _set_ind(pc, leftChars=200, left=544, hangingChars=100, hanging=184)
    # 解説（前文・構造式・選択肢解説・解答）。見出しは上で付けたので add_header=False
    append_explanation(doc, sub, item, add_header=False)


def build_from_pdf_result(result):
    """PDF応答（{'小問':[...]}）から house形式の docx を生成する。"""
    items = result.get('小問', []) or []
    doc = new_house_doc()
    if len(items) > 1:
        nums = [str(it.get('問番号', '')) for it in items]
        t = doc.add_paragraph()
        t._p.append(make_run('問' + nums[0] + '〜' + nums[-1], font=GOTHIC, bold=True))
    for it in items:
        append_pdf_item(doc, it)
    return doc, items


SYSTEM_PROMPT = """あなたは薬剤師国家試験の解説作成の専門家です。
以下のJSON形式のみで回答してください（他のテキスト不要）:
{
  "前文": "概念整理・公式・前提（純粋な知識問題で不要なら空文字）",
  "前文_追加行": ["前文が複数段落になる場合の2段落目", "3段落目"],
  "構造式": [{"ラベル":"アミド", "smiles":"CC(=O)N"}],
  "選択肢解説": [
    {"番号":"１","正誤":"誤","内容":"解説文"},
    {"番号":"２","正誤":"正","内容":"解説文"},
    {"番号":"３","正誤":"誤","内容":"解説文"},
    {"番号":"４","正誤":"誤","内容":"解説文"},
    {"番号":"５","正誤":"正","内容":"解説文"}
  ],
  "解答": ["２","５"]
}

【文体・スタイルのルール】
- 文末：「〜である。」「〜となる。」「〜と考えられる。」「〜と推定できる。」など断定調。「〜です」「〜ます」は使わない。
- 「誤」の選択肢：なぜ誤りかを説明し、「なお、正しくは〜である。」や「〜ではなく、〜である。」で正しい情報も必ず示す。
- 前文で説明済みの内容は「上記参照。」でよい。
- 他の選択肢を参照するときは「選択肢X参照。」と書く。
- 計算問題では前文に「●Step1：」「●Step2：」などの段階的な解法を示す。
- 薬学専門用語・薬物名・受容体名は正確に記載する。
- 計算は数値を具体的に示し、途中式を省略しない。

【記号・表記の統一ルール】
- ハイフンとマイナス：数式・モデル名・化合物名・略号内の連結記号にはすべて「－」（全角ハイフンマイナス、U+FF0D）を使用し、半角ハイフン「-」は一切使わない。
  例：「1－コンパートメントモデル」「P－糖タンパク質」「P－gp」「UDP－グルクロン酸転移酵素」「L－カルボシステイン」「α－ヘリックス」（「1-」「P-」のような半角ハイフンは禁止）
- 引き算の記号はすべて「－」（全角マイナス）。例：50－30＝20
- クリアランスの区別：
  ・CL = 全身クリアランス（肝・腎などをすべて含む）
  ・CLr = 腎クリアランス（renal clearance）
  ・CLh = 肝クリアランス（hepatic clearance）
  ・問題文でCLtotと明示されている場合のみCLtotを使用
- 投与量：D（一般）、Dpo（経口）、Div（静注）。"Dose"は使わない。
- 消失速度定数：kel（keは使わない）
- 吸収速度定数：ka
- 分布容積：Vd
- バイオアベイラビリティ：F
- 最高血中濃度：Cmax
- 定常状態血中濃度：Css（平均はCss,av）
- 消失半減期：t1/2（T1/2は使わない）
- 投与間隔：τ
- 平均滞留時間：MRT、平均吸収時間：MAT、平均溶出時間：MDT
- 式番号：①②③…を用いて文中で相互参照
- 単位：h、L、mg など（h^{-1} のように指数はマークアップで表す）

【上付き・下付き文字のルール（重要）】
- 下付きは _{ } 、上付きは ^{ } で必ず明示する（LaTeX風）。地の文にそのまま数字を並べない。
  例：初期濃度 A_{0}、半減期 t_{1/2}、定常状態 C_{ss}、最高濃度 C_{max}
  例：指数・べき乗 e^{-kt}、1×10^{-5}、10^{6}、速度定数の単位 h^{-1}、面積 m^{2}
- 化学式は、原子数を下付き、電荷を上付きで表す。
  例：MnO_{4}^{-}、C_{2}O_{4}^{2-}、Mn^{2+}、CO_{2}、H^{+}、SO_{4}^{2-}、Ca^{2+}
- 反応式もこの表記で書く。例：2MnO_{4}^{-}＋5C_{2}O_{4}^{2-}＋16H^{+}→2Mn^{2+}＋10CO_{2}＋8H_{2}O
- 「1次反応」「2倍」など、数字が下付き・上付きでない箇所には付けない（通常の全角数字のまま）。

【構造式（化学構造）の提示】
- 有機化学・医薬品化学など、構造を示すと理解が深まる問題では "構造式" フィールドに
  各分子を {"ラベル":"名称", "smiles":"SMILES文字列"} の形で列挙する。前文の後に図版として挿入される。
- SMILESは正確なもののみ記載する（自信がない構造は載せない）。ラベルは日本語名や記号（A・B・C 等）でよい。
  例）カルボン酸誘導体の反応性比較：
  "構造式":[{"ラベル":"酸ハロゲン化物","smiles":"CC(=O)Cl"},{"ラベル":"酸無水物","smiles":"CC(=O)OC(C)=O"},
           {"ラベル":"エステル","smiles":"CC(=O)OC"},{"ラベル":"カルボン酸","smiles":"CC(=O)O"},
           {"ラベル":"アミド","smiles":"CC(=O)N"}]
- 構造が不要な問題では "構造式" は省略（空配列でよい）。

【複数の場合・種類を列挙するときのレイアウト】
- 前文で複数の場合（例：0次反応・1次反応・2次反応、酸性/中性/塩基性 など）を提示するときは、
  各項目を必ず「●」で始め、"前文_追加行" を使って1項目=1行に分けて記載する（1段落に詰め込まない）。
  例）前文＝「反応次数ごとに速度式と半減期を整理する。」
      前文_追加行＝[
        "●0次反応：A＝A_{0}－kt。半減期 t_{1/2}＝A_{0}／(2k) で A_{0} に比例する。",
        "●1次反応：A＝A_{0}e^{-kt}。半減期 t_{1/2}＝0.693／k で A_{0} によらず一定。",
        "●2次反応：1／A＝1／A_{0}＋kt。半減期 t_{1/2}＝1／(k・A_{0}) で A_{0} に反比例する。"
      ]

【良い解説の例①：薬剤（薬物動態）知識問題（第110回 問173）】
正解：３、５
前文：「薬物の尿中排泄は、糸球体におけるろ過、尿細管における分泌、再吸収という三つの過程によって行われる。血液中に含まれる薬物のうちタンパクと結合していない非結合形の薬物が糸球体でろ過を受ける。次に、近位尿細管においてトランスポーター等によって認識される薬物は分泌を受け、尿細管に流入する。その後、尿細管内に流入した薬物のうち再吸収を受けなかった薬物が尿中に排泄される。このうち、糸球体ろ過における単位時間あたりの血漿のろ過量を糸球体ろ過速度（GFR）といい、通常成人では約100 mL/min/1.73 m2である。ろ過クリアランスは、GFR×fp（fp：血漿タンパク非結合率）で表される。イヌリンは血漿タンパクと結合せず（fp=1）、尿細管分泌や再吸収を受けないため、イヌリンクリアランス＝GFRとなる。一方、クレアチニンは尿細管において若干の分泌を受けるため、クレアチニンクリアランス＞GFRとなる。」
選択肢１：「誤：GFR＝イヌリンクリアランス＝30 mL/min/1.73 m2と推定できる。」
選択肢２：「誤：イヌリンは尿細管で分泌・再吸収を受けないため、再吸収クリアランスは存在しない。」
選択肢３：「正：クレアチニンの尿細管分泌クリアランス＝クレアチニンクリアランス－イヌリンクリアランス＝50－30＝20 mL/min/1.73 m2と推定できる。」
選択肢４：「誤：正常成人のGFRは約100 mL/min/1.73 m2であるため、本患者のイヌリンクリアランス30 mL/min/1.73 m2は正常時より小さいと考えられる。」
選択肢５：「正：本患者のクレアチニンクリアランス50 mL/min/1.73 m2は正常時（約100 mL/min/1.73 m2）より小さいと考えられる。」

【良い解説の例②：薬理知識問題（第110回 問160）】
正解：１、５
前文：「」（純粋知識問題のため前文なし）
選択肢１：「正：オキシメテバノールは麻薬性鎮咳薬であり、オピオイド受容体を刺激して鎮咳作用を示す。」
選択肢２：「誤：L－カルボシステインは構造中にSH基を有さず、ムコタンパク質の構成成分であるシアル酸・フコースのバランスを改善して去痰作用を示す。なお、SH基を有しムコタンパク質のペプチド鎖の連結を切断して去痰作用を示す薬剤はアセチルシステインである。」
選択肢３：「誤：フルマゼニルはベンゾジアゼピン受容体に結合し、ベンゾジアゼピン系薬の作用に拮抗することで呼吸抑制を改善する。なお、末梢性化学受容器を刺激して間接的に呼吸中枢を興奮させる薬剤はドキサプラムである。」
選択肢４：「誤：アンブロキソールはブロムヘキシンの活性代謝物であり、肺サーファクタント分泌の促進・線毛運動の亢進により去痰作用を示す。（本選択肢は親薬物と活性代謝物が逆）」
選択肢５：「正：ニンテダニブは低分子チロシンキナーゼ阻害薬であり、VEGFR・FGFR・PDGFRのチロシンキナーゼを阻害して肺の線維化を抑制する。」

【良い解説の例③：生化学知識問題（第110回 問115）】
正解：２、３
前文：「」（純粋知識問題のため前文なし）
選択肢１：「誤：α－ヘリックスとは1本のポリペプチド鎖が分子内水素結合することで形成される細長いらせん構造を指す。コラーゲンの三重らせん構造は3本のポリペプチド鎖が形成する特殊ならせん構造であり、α－ヘリックスとは異なる。」
選択肢２：「正：Xはプロリン（Pro）の翻訳後修飾で生じたアミノ酸であることから、ヒドロキシプロリンであると考えられる。」
選択肢３：「正：グリシン（Gly）は側鎖が水素原子であるため空間を占める割合が小さく、コラーゲンの三重らせん構造における混み合った部位に安定して収まることができるため、三重らせん構造の形成に重要な役割を担う。」
選択肢４：「誤：ビタミンCがコラーゲン遺伝子の転写を促進する際に核内に移行するかどうかは現在不明である。」
選択肢５：「誤：コラーゲンは細胞外マトリックスの代表的な構成タンパク質であり組織の強度を保つが、細胞内の細胞骨格を構成するわけではない。」

【良い解説の例④：グラフ問題（第111回 問170）】
正解：１、５
前文：「リボフラビンは、食事の有無による胃内容排出速度（GER）の違いにより吸収量が変化する。空腹時に服用するとGERの増大により、十二指腸に存在する吸収トランスポーターが飽和しやすくなり、吸収量が低下する。一方、食後に服用するとGERの低下により、吸収トランスポーターの飽和が起こりにくくなることで吸収量が増大する。なお、グラフの縦軸は累積尿中排泄量を示しているが、累積尿中排泄量≒吸収量と考える必要があるため、Aが朝食後服用（吸収量が多い）、Bが空腹時服用（吸収量が少ない）と読み取れる。」
選択肢１：「正：上記参照。」
選択肢２：「誤：上記参照。」
選択肢３：「誤：BがAより低値となるのは、リボフラビンのGERが空腹時に増大することによる吸収トランスポーターの飽和が原因である。」
選択肢４：「誤：選択肢３参照。」
選択肢５：「正：メトクロプラミドはドパミンD2受容体遮断薬であり、D2受容体を遮断することでコリン作動性神経を興奮させGERを増大させる。そのため、メトクロプラミドを前投与した時の曲線は空腹時服用時と類似し、AよりBに近くなる。」

マークアップ:
- 下付き _{...} / 上付き ^{...}（例：A_{0}、t_{1/2}、10^{-5}、MnO_{4}^{-}）
- 特殊フォントの変数：{{CL:r}}=CLr {{f:e}}=fe {{K:sp}}=Ksp {{Vd}}=Vd {{mu}}=µ
- ハイフン・マイナス（連結記号）は全角「－」を使用"""


RENMON_INSTRUCTION = """
【連問（問196以降の実践問題）の指示】
これは複数の小問が連続する「連問」である。以下の小問すべての解説を、必ず次のJSON形式（連問配列）のみで一括生成すること:
{{
  "連問": [
    {{"問番号":"{first}", "前文":"...", "前文_追加行":[], "選択肢解説":[{{"番号":"１","正誤":"誤","内容":"..."}}, ...], "解答":["３"]}},
    ...（小問の数だけ続ける）...
  ]
}}
対象の小問（この順序・この問番号で出力すること）: {targets}
- 連問は前の小問の内容を前提とする（例：「前問の処方提案の理由は〜」）。共有の症例・検査値・処方を踏まえ、各小問の整合性を保つこと。
- 各小問の「選択肢解説」は、その小問の選択肢数に一致させること。
- 「1つ選べ」の小問は解答を1つ、「2つ選べ」は2つ、という具合に設問の指示に従うこと。"""


PDF_INSTRUCTION = """添付のPDFは薬剤師国家試験の問題です（図・グラフ・構造式を含むことがあります）。
PDF内の図や構造式もよく見て内容を理解し、各小問の解説を作成してください。

必ず次のJSON形式のみで回答してください（他のテキスト不要）:
{
  "小問": [
    {
      "問番号": "102",
      "科目": "物理・化学・生物",
      "設問": "設問文（〜はどれか。1つ選べ。 など）",
      "選択肢": ["選択肢1の内容（構造なら化合物名）", "選択肢2の内容", "..."],
      "前文": "概念整理・前提（不要なら空文字）",
      "前文_追加行": [],
      "構造式": [{"ラベル":"名称", "smiles":"SMILES"}],
      "選択肢解説": [{"番号":"１","正誤":"誤","内容":"..."}, "..."],
      "解答": ["４"]
    }
  ]
}

【重要】
- PDFの図・構造式を実際に読み取り、選択肢と構造の対応・正誤を正確に判定すること（推測で番号を取り違えない）。
- 連問（問196以降など複数小問）はPDF内の小問をすべて "小問" 配列に含め、前後の整合性を保つこと。
- 構造が関わる問題は "構造式" に登場分子のSMILESを列挙する（解説に構造図として挿入される）。
- 文体・記号（_{ }/^{ } の上付き下付き、全角ハイフン、●列挙など）は上記の共通ルールに従うこと。"""


def _extract_json(text):
    depth, start = 0, None
    for i, ch in enumerate(text):
        if ch == '{':
            if depth == 0:
                start = i
            depth += 1
        elif ch == '}':
            depth -= 1
            if depth == 0 and start is not None:
                return json.loads(text[start:i + 1])
    raise ValueError("JSONが見つかりません")


def _post_api(payload, api_key):
    req = urllib.request.Request(
        "https://api.anthropic.com/v1/messages",
        data=json.dumps(payload).encode(),
        headers={'Content-Type': 'application/json',
                 'x-api-key': api_key,
                 'anthropic-version': '2023-06-01'},
        method='POST')
    with urllib.request.urlopen(req, timeout=240) as resp:
        text = json.loads(resp.read())['content'][0]['text']
        return _extract_json(text)


def call_api(question_text, api_key, sub_questions=None):
    """単問なら解説JSONを、連問なら {'連問':[...]} を返す。
    sub_questions: 連問時の [{'番号':..,'科目':..,'選択肢数':..}] のリスト。"""
    user_content = f"解説を生成してください:\n\n{question_text}"
    if sub_questions and len(sub_questions) > 1:
        targets = "、".join(
            f"問{s['番号']}（{s.get('科目') or ''}・選択肢{s.get('選択肢数', 5)}）"
            for s in sub_questions)
        user_content += RENMON_INSTRUCTION.format(
            first=sub_questions[0]['番号'], targets=targets)
    payload = {
        "model": "claude-opus-4-8",
        "max_tokens": 8192,
        "system": SYSTEM_PROMPT,
        "messages": [{"role": "user", "content": user_content}]
    }
    return _post_api(payload, api_key)


def call_api_pdf(pdf_bytes, api_key):
    """PDF（Word発。図・構造式が描画済み）を画像認識APIに渡し、{'小問':[...]} を返す。"""
    import base64
    b64 = base64.b64encode(pdf_bytes).decode()
    content = [
        {"type": "document",
         "source": {"type": "base64", "media_type": "application/pdf", "data": b64}},
        {"type": "text", "text": PDF_INSTRUCTION},
    ]
    payload = {
        "model": "claude-opus-4-8",
        "max_tokens": 8192,
        "system": SYSTEM_PROMPT,
        "messages": [{"role": "user", "content": content}]
    }
    return _post_api(payload, api_key)


PDF_EXPLAIN_INSTRUCTION = """添付PDFはこの問題（図・グラフ・構造式を含むことがある）です。
図・構造式もよく見て理解し、次の各小問について「解説だけ」をJSONで返してください。
設問文・選択肢の再掲は不要です（それらは別のWordテンプレート側にあります）。

必ず次のJSON形式のみで回答（他のテキスト不要）:
{{
  "小問": [
    {{"問番号":"{first}", "前文":"...", "前文_追加行":[],
      "構造式":[{{"ラベル":"名称","smiles":"SMILES"}}],
      "選択肢解説":[{{"番号":"１","正誤":"誤","内容":"..."}}, "..."], "解答":["４"]}}
  ]
}}
対象の小問（この順序・この問番号で出力）: {targets}

【重要】
- PDFの図・構造式を実際に読み取り、選択肢と構造の対応・正誤を正確に判定すること。
  （例：O−アシルイソ尿素・混合酸無水物などの活性化アシル種は反応性が高い。求核アシル置換の
   反応性序列＝酸ハロゲン化物＞酸無水物・活性エステル＞エステル＞アミド を丁寧に当てはめる。）
- "構造式" には「解説で新たに図示すべき構造」だけをSMILESで入れる
  （反応の生成物・活性体・代謝物・中間体など）。問題に既に描かれている構造は入れない。
  新規構造が不要なら "構造式" は空配列でよい。
- 文体・記号（下付き _{{ }}／上付き ^{{ }}、全角ハイフン、●列挙）は共通ルールに従う。"""


def call_api_pdf_explain(pdf_bytes, api_key, subs):
    """PDF（図・構造式つき）を見せて、docxの各小問に対応する「解説だけ」を {'小問':[...]} で返す。"""
    import base64
    b64 = base64.b64encode(pdf_bytes).decode()
    targets = "、".join(
        f"問{s['num']}（{s.get('subject') or ''}・選択肢{s.get('n_choices', 5)}）" for s in subs)
    instr = PDF_EXPLAIN_INSTRUCTION.format(first=subs[0]['num'], targets=targets)
    content = [
        {"type": "document",
         "source": {"type": "base64", "media_type": "application/pdf", "data": b64}},
        {"type": "text", "text": instr},
    ]
    payload = {
        "model": "claude-opus-4-8",
        "max_tokens": 8192,
        "system": SYSTEM_PROMPT,
        "messages": [{"role": "user", "content": content}]
    }
    return _post_api(payload, api_key)


# ──────────────────────────────────────────────
# Streamlit UI
# ──────────────────────────────────────────────
def main():
    st.set_page_config(
        page_title='薬剤師国家試験 解説生成システム',
        page_icon='📋',
        layout='centered'
    )

    st.title('📋 薬剤師国家試験 解説生成システム')
    st.markdown('テンプレートdocxをアップロードして解説を自動生成します。')
    st.divider()

    # ① ファイルアップロード
    st.subheader('① ファイルをアップロード')
    uploaded = st.file_uploader(
        'docx（テンプレート）または PDF（図・構造式つき問題）をドラッグ＆ドロップ',
        type=['docx', 'pdf'],
        help='docx＝問題と「ああ」プレースホルダー入りテンプレート。'
             'PDF＝Wordで「PDFとして書き出し」した問題（構造式・グラフが描画されているもの）。'
             'PDFは画像認識で図を読み取り、解説Word（構造式つき）を新規生成します。'
    )
    is_pdf = bool(uploaded) and uploaded.name.lower().endswith('.pdf')
    is_docx = bool(uploaded) and uploaded.name.lower().endswith('.docx')
    vision_pdf = None
    if is_pdf:
        st.info('📄 PDF単体モード：図・構造式を画像認識で読み取り、解説Wordを新規生成します（AIモードのみ）。')
    if is_docx:
        vision_pdf = st.file_uploader(
            '（任意）構造式・グラフのある問題は、その問題のPDF（Wordで「PDF書き出し」したもの）も添付',
            type=['pdf'], key='vision_pdf',
            help='PDFを添付すると、AIが図・構造式を実際に見て解説を書きます。'
                 '出力はこのdocx原本のまま（問題・構造・体裁は保持）、解説だけ差し込みます。')
        if vision_pdf is not None:
            st.info('🖼️ docx＋PDFモード：docx原本はそのまま、AIがPDFの図を見て解説を差し込みます。')

    # ② モード
    st.subheader('② モードを選択')
    mode = st.radio('', ['🤖 AIが自動生成', '✏️ 手動入力（無料）'],
                    horizontal=True, label_visibility='collapsed')

    # ③ 入力
    st.subheader('③ 解説内容')
    explanation = None

    if 'AI' in mode:
        # APIキー：Streamlit secretsにあれば自動使用
        if 'ANTHROPIC_API_KEY' in st.secrets:
            api_key = st.secrets['ANTHROPIC_API_KEY']
            st.success('✅ APIキー設定済み（管理者設定）')
        else:
            api_key = st.text_input(
                'Anthropic APIキー',
                type='password',
                placeholder='sk-ant-api03-...',
                help='console.anthropic.com で取得。1問あたり約1〜2円。'
            )
    else:
        api_key = None
        st.markdown('各選択肢の解説を入力してください。')
        st.caption('マークアップ: `{{CL:r}}`=CLr　`{{f:e}}`=fe　`{{-}}`=全角マイナス　`{{mu}}`=µ')

        manual_data = {}
        for char in NUMBERS:
            col1, col2 = st.columns([1, 5])
            with col1:
                seigo = st.selectbox(char, ['誤', '正'], key=f'seigo_{char}', label_visibility='visible')
            with col2:
                content = st.text_input(f'選択肢{char}の解説', key=f'content_{char}',
                                        label_visibility='collapsed',
                                        placeholder=f'選択肢{char}の解説を入力...')
            manual_data[char] = {'seigo': seigo, 'content': content}

        st.markdown('')
        col_a, col_b = st.columns(2)
        with col_a:
            ans_input = st.text_input('解答（例: ２,５）', placeholder='２,５')
        with col_b:
            preamble = st.text_input('前文（任意）', placeholder='計算式など')

    st.divider()

    # 生成ボタン
    generate = st.button('🚀 解説を生成する', type='primary', use_container_width=True)

    if generate:
        # バリデーション
        if not uploaded:
            st.error('ファイルをアップロードしてください。')
            st.stop()

        if 'AI' in mode and not api_key:
            st.error('APIキーを入力してください。')
            st.stop()

        # ===== PDFモード（画像認識）: 図・構造式を読み取り解説Wordを新規生成 =====
        if is_pdf:
            if '手動' in mode:
                st.error('PDFモードはAIモードのみ対応です。上で「AIが自動生成」を選んでください。')
                st.stop()
            with st.spinner('PDFの図・構造式を読み取り、解説を生成中...（少し時間がかかります）'):
                try:
                    pdf_bytes = uploaded.read()
                    result = call_api_pdf(pdf_bytes, api_key)
                    out_doc, items = build_from_pdf_result(result)
                    if not items:
                        st.error('PDFから小問を読み取れませんでした。ファイルをご確認ください。')
                        st.stop()

                    buf = io.BytesIO()
                    out_doc.save(buf)
                    buf.seek(0)
                    fname = uploaded.name.rsplit('.', 1)[0] + '_解説.docx'
                    ans_str = ' ／ '.join(
                        f"問{it.get('問番号', '')}: " + '、'.join(it.get('解答', []) or [])
                        for it in items)
                    n_struct = sum(len(it.get('構造式', []) or []) for it in items)
                    st.success(f'✅ 完成！（解答 → {ans_str}）')
                    if n_struct:
                        st.caption(f'構造式データ {n_struct} 件を受信（RDKitで描画を試行しました）。')
                    else:
                        st.caption('（この問題ではAIから構造式データは出力されませんでした）')
                    st.download_button(
                        label='📥 解説Wordをダウンロード',
                        data=buf,
                        file_name=fname,
                        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        use_container_width=True)
                except urllib.error.HTTPError as e:
                    st.error(f'APIエラー（{e.code}）: APIキー・PDF形式をご確認ください。')
                except Exception as e:
                    st.error(f'エラーが発生しました: {e}')
                    st.exception(e)
            st.stop()

        # ===== docx＋PDFモード：docx原本はそのまま、PDFで図を見て解説を差し込む =====
        if is_docx and vision_pdf is not None:
            if '手動' in mode:
                st.error('docx＋PDFモードはAIモードのみ対応です。')
                st.stop()
            with st.spinner('PDFの図・構造式を読み取り、docx原本に解説を差し込み中...'):
                try:
                    doc = Document(io.BytesIO(uploaded.read()))
                    subs, boundary = parse_question_blocks(doc)
                    if not subs:
                        st.error('docxの問題文から小問（問◯◯）を検出できませんでした。')
                        st.stop()
                    st.info('・'.join(f"問{s['num']}" for s in subs) + ' を検出。PDFの図を読み取り中...')

                    result = call_api_pdf_explain(vision_pdf.read(), api_key, subs)
                    items = result.get('小問', []) if isinstance(result, dict) else []
                    if not items:
                        raise ValueError('PDFから解説（小問配列）が得られませんでした。')
                    by_num = {str(it.get('問番号')): it for it in items}
                    expls = []
                    for i, s in enumerate(subs):
                        it = by_num.get(str(s['num']))
                        if it is None:
                            it = items[i] if i < len(items) else {}
                        expls.append(it)

                    regenerate_explanations(doc, subs, boundary, expls)

                    buf = io.BytesIO()
                    doc.save(buf)
                    buf.seek(0)
                    fname = uploaded.name.replace('.docx', '_完成.docx')
                    ans_str = ' ／ '.join(
                        f"問{s['num']}: " + '、'.join(e.get('解答', []) or [])
                        for s, e in zip(subs, expls))
                    n_struct = sum(len(e.get('構造式', []) or []) for e in expls)
                    st.success(f'✅ 完成！（解答 → {ans_str}）')
                    if n_struct:
                        st.caption(f'新規構造式データ {n_struct} 件を受信（RDKitで描画を試行しました）。')
                    st.download_button(
                        label='📥 完成ファイルをダウンロード',
                        data=buf,
                        file_name=fname,
                        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        use_container_width=True)
                except urllib.error.HTTPError as e:
                    st.error(f'APIエラー（{e.code}）: APIキー・PDF形式をご確認ください。')
                except Exception as e:
                    st.error(f'エラーが発生しました: {e}')
                    st.exception(e)
            st.stop()

        if '手動' in mode:
            missing = [c for c in NUMBERS if not manual_data[c]['content'].strip()]
            if missing:
                st.error(f'選択肢{"、".join(missing)}の解説を入力してください。')
                st.stop()
            if not ans_input.strip():
                st.error('解答を入力してください。')
                st.stop()

            explanation = {
                '前文': preamble.strip() if preamble else '',
                '前文_追加行': [],
                '選択肢解説': [
                    {'番号': c, '正誤': manual_data[c]['seigo'], '内容': manual_data[c]['content']}
                    for c in NUMBERS
                ],
                '解答': [a.strip() for a in ans_input.split(',')]
            }

        # 処理
        with st.spinner('処理中...'):
            try:
                doc = Document(io.BytesIO(uploaded.read()))
                # 問題本文から実際の小問を自動検出（社内フォーマットの雛形はそのままでよい）
                subs, boundary = parse_question_blocks(doc)

                if not subs:
                    st.error('問題文から小問（問◯◯）を検出できませんでした。ファイルを確認してください。')
                    st.stop()

                is_renmon = len(subs) > 1
                q_labels = '・'.join(f"問{s['num']}" for s in subs)
                if is_renmon:
                    st.info(f'連問を検出しました（{q_labels}／{len(subs)}問）')
                else:
                    st.info(f'問{subs[0]["num"]} を検出しました')

                if 'AI' in mode:
                    # 問題本文のみ抽出（boundary=雛形開始より前）＋表
                    lines = [p.text.strip() for p in doc.paragraphs[:boundary] if p.text.strip()]
                    for table in doc.tables:
                        lines.append('')
                        for row in table.rows:
                            cells = [c.text.replace('\n', ' ').strip() for c in row.cells]
                            if any(cells):
                                lines.append(' | '.join(cells))
                    question_text = '\n'.join(lines)

                    st.info('Claude API で解説生成中...')
                    sub_meta = [{'番号': s['num'], '科目': s.get('subject'),
                                 '選択肢数': s.get('n_choices', 5)} for s in subs]
                    result = call_api(question_text, api_key,
                                      sub_questions=sub_meta if is_renmon else None)

                    if is_renmon:
                        items = result.get('連問', []) if isinstance(result, dict) else []
                        if not items:
                            raise ValueError('連問形式の応答が得られませんでした。')
                        by_num = {str(it.get('問番号')): it for it in items}
                        expls = []
                        for i, s in enumerate(subs):
                            it = by_num.get(str(s['num']))
                            if it is None:
                                it = items[i] if i < len(items) else {}
                            expls.append(it)
                    else:
                        expls = [result]
                else:
                    if is_renmon:
                        st.error('連問は手動入力に未対応です。AIモードをご利用ください。')
                        st.stop()
                    expls = [explanation]

                # 雛形を削除し、実番号で解説ブロックを再生成
                regenerate_explanations(doc, subs, boundary, expls)

                # docxをメモリに保存
                buf = io.BytesIO()
                doc.save(buf)
                buf.seek(0)

                fname = uploaded.name.replace('.docx', '_完成.docx')
                ans_str = ' ／ '.join(
                    f"問{s['num']}: " + '、'.join(e.get('解答', []))
                    for s, e in zip(subs, expls))

                st.success(f'✅ 完成！（解答 → {ans_str}）')

                st.download_button(
                    label='📥 完成ファイルをダウンロード',
                    data=buf,
                    file_name=fname,
                    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    use_container_width=True
                )

            except urllib.error.HTTPError as e:
                st.error(f'APIエラー（{e.code}）: APIキーを確認してください。')
            except Exception as e:
                st.error(f'エラーが発生しました: {e}')
                st.exception(e)


if __name__ == '__main__':
    main()
